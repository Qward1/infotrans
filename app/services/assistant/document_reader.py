"""Чтение документов (PDF/DOCX/TXT/Markdown) и сохранение как ``Document``.

Извлечённый текст связывается с пользователем (и опционально с событием) и
используется генератором протоколов. Если библиотека для формата недоступна —
возвращаем понятное предупреждение, не роняя запрос.
"""
from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field

from sqlalchemy.orm import Session

from app.models.assistant import Document

logger = logging.getLogger("smartcal.documents")

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


@dataclass
class ExtractedDocument:
    filename: str
    content_type: str
    size_bytes: int
    text: str = ""
    warnings: list[str] = field(default_factory=list)


def _ext(filename: str) -> str:
    return "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def _read_pdf(raw: bytes, doc: ExtractedDocument) -> None:
    try:
        from pypdf import PdfReader
    except ImportError:
        doc.warnings.append("Не установлен pypdf — не могу прочитать PDF (pip install pypdf).")
        return
    try:
        reader = PdfReader(io.BytesIO(raw))
        parts = [(page.extract_text() or "") for page in reader.pages]
        doc.text = "\n".join(parts).strip()
        if not doc.text:
            doc.warnings.append("PDF без текстового слоя (возможно, скан) — текст не извлечён.")
    except Exception as exc:  # noqa: BLE001
        doc.warnings.append(f"Ошибка чтения PDF: {exc}")


def _read_docx(raw: bytes, doc: ExtractedDocument) -> None:
    try:
        import docx  # python-docx
    except ImportError:
        doc.warnings.append("Не установлен python-docx — не могу прочитать DOCX.")
        return
    try:
        d = docx.Document(io.BytesIO(raw))
        lines = [p.text for p in d.paragraphs]
        # таблицы тоже часто содержат решения/задачи
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        doc.text = "\n".join(l for l in lines if l is not None).strip()
    except Exception as exc:  # noqa: BLE001
        doc.warnings.append(f"Ошибка чтения DOCX: {exc}")


def extract(filename: str, content_type: str, raw: bytes) -> ExtractedDocument:
    """Извлечь текст из документа по расширению."""
    doc = ExtractedDocument(
        filename=filename or "document",
        content_type=content_type or "application/octet-stream",
        size_bytes=len(raw),
    )
    ext = _ext(filename or "")
    if ext in {".txt", ".md"}:
        doc.text = raw.decode("utf-8", errors="replace").strip()
    elif ext == ".pdf":
        _read_pdf(raw, doc)
    elif ext == ".docx":
        _read_docx(raw, doc)
    else:
        doc.warnings.append(f"Неподдерживаемый тип файла: {ext or 'без расширения'}")
    return doc


def save(
    db: Session,
    owner_id: int,
    extracted: ExtractedDocument,
    event_id: int | None = None,
) -> Document:
    """Сохранить извлечённый документ в БД."""
    row = Document(
        owner_id=owner_id,
        event_id=event_id,
        filename=extracted.filename,
        content_type=extracted.content_type,
        size_bytes=extracted.size_bytes,
        text=extracted.text,
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    return row


def get(db: Session, document_id: int) -> Document | None:
    return db.get(Document, document_id)


def delete(db: Session, document: Document) -> None:
    """Удалить документ (FN-12). Протоколы — payload действий, ссылок на строку нет."""
    db.delete(document)
    db.commit()


def list_for_user(db: Session, owner_id: int, limit: int = 50) -> list[Document]:
    """Документы пользователя, новые сверху."""
    from sqlalchemy import select

    stmt = (
        select(Document)
        .where(Document.owner_id == owner_id)
        .order_by(Document.created_at.desc(), Document.id.desc())
        .limit(limit)
    )
    return list(db.execute(stmt).scalars().all())
